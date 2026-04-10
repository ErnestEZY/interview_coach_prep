from typing import Tuple
from docx import Document
from pdfminer.high_level import extract_text
from pypdf import PdfReader
import pdfplumber
import pytesseract
from pdf2image import convert_from_path
import os
import re

def is_pdf(filename: str) -> bool:
    return filename.lower().endswith(".pdf")

def is_docx(filename: str) -> bool:
    return filename.lower().endswith(".docx") or filename.lower().endswith(".doc")

def clean_text(text: str) -> str:
    """Removes extra whitespace and normalize text."""
    if not text:
        return ""
    # Replace multiple newlines with single one
    text = re.sub(r'\n+', '\n', text)
    # Replace multiple spaces with single one
    text = re.sub(r' +', ' ', text)
    return text.strip()

def extract_resume_text(path: str) -> Tuple[str, str]:
    name = os.path.basename(path)
    text = ""
    mime = ""

    if is_pdf(name):
        mime = "application/pdf"
        
        # 1. Try pypdf (Fastest and usually sufficient for text-based PDFs)
        try:
            reader = PdfReader(path)
            pypdf_text = ""
            for page in reader.pages:
                pypdf_text += (page.extract_text() or "") + "\n"
            text = pypdf_text.strip()
        except Exception:
            text = ""

        # 2. Try pdfplumber (Better for complex layouts/tables, but slower)
        if len(text) < 300:
            try:
                with pdfplumber.open(path) as pdf:
                    plumber_text = ""
                    for page in pdf.pages:
                        plumber_text += (page.extract_text() or "") + "\n"
                    text = plumber_text.strip()
            except Exception:
                pass

        # 3. Fallback to pdfminer
        if len(text) < 100:
            try:
                text = extract_text(path).strip()
            except Exception:
                pass

        # Check if PDF is likely image-based (no selectable text)
        if not text or len(text) < 100:
            raise ValueError(
                "This PDF appears to be a scanned image or lacks selectable text. "
                "Image-based PDFs are not compatible with ATS systems. "
                "Please upload a standard PDF with selectable text, or a Word (.docx) file."
            )

        # 4. Optimized OCR (Only for borderline cases to keep speed high)
        if len(text) < 800: # Only run OCR if text is very sparse
            try:
                with pdfplumber.open(path) as pdf:
                    ocr_additions = ""
                    for page in pdf.pages:
                        # Limit images per page to top 3 largest to save time
                        page_images = sorted(page.images, key=lambda x: (x["x1"]-x["x0"])*(x["bottom"]-x["top"]), reverse=True)[:3]
                        
                        for img in page_images:
                            try:
                                # Skip tiny images (likely icons/bullets) that don't contain meaningful text
                                width = img["x1"] - img["x0"]
                                height = img["bottom"] - img["top"]
                                if width < 30 or height < 10:
                                    continue
                                    
                                bbox = (img["x0"], img["top"], img["x1"], img["bottom"])
                                cropped = page.within_bbox(bbox).to_image(resolution=150) # Lower resolution for speed
                                img_text = pytesseract.image_to_string(cropped.original, config='--psm 6').strip()
                                if img_text and len(img_text) > 2:
                                    ocr_additions += "\n" + img_text
                            except Exception:
                                continue
                    if ocr_additions:
                        text += "\n" + ocr_additions
                        print(f"Captured additional visual text: {len(ocr_additions)} characters.")
            except Exception as e:
                print(f"Visual OCR extraction failed: {e}")

        return clean_text(text), mime

    if name.lower().endswith(".doc"):
        raise ValueError("Please convert .doc to .docx or pdf")
    
    if is_docx(name):
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        doc = Document(path)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        # Extract text from tables in Word
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
        
        if not text.strip():
             raise ValueError("The Word document appears to be empty.")
        return clean_text(text), mime
    
    raise ValueError(f"Unsupported file type: {name}. Please upload a PDF or DOCX file.")
