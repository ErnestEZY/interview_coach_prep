"""Integration Tests — Resume API"""
import os, sys, io, pytest, pytest_asyncio
from unittest.mock import AsyncMock, patch
from httpx import AsyncClient, ASGITransport
from datetime import datetime, timezone
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
if ROOT not in sys.path: sys.path.insert(0, ROOT)

from tests.integration.helpers import make_jwt, patch_all_db

UID = "507f191e810c19729de860ea"
BASE_USER = {"_id": UID, "email": "u@t.com", "role": "user", "is_verified": True,
             "has_analyzed": False, "daily_resume_count": 0,
             "daily_interview_count": 0, "daily_question_count": 0, "daily_reset_at": None}
FAKE_FB = {
    "IsResume": True, "Score": 75,
    "ScoreBreakdown": {"ImpactScore": 30, "SkillScore": 22, "StructureScore": 16, "ATSScore": 7},
    "Advantages": ["Good"], "Disadvantages": ["Weak"], "Suggestions": ["Add summary"],
    "Keywords": ["Python"], "Location": "KL", "DetectedJobTitle": "Software Engineer",
    "Email": "", "Phone": "", "Website": "", "ProfessionalSummary": "Dev.",
    "Education": [], "Experience": [], "Projects": [],
    "SkillsTech": "Python", "SkillsTools": "Git", "SkillsSoft": "Comm",
    "Certifications": [], "Languages": [], "AdditionalInfo": [],
}

def _docx():
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Jane Doe | jane@test.com | +60123456789")
    doc.add_paragraph("Summary: Python developer, 3 years experience in FastAPI.")
    doc.add_paragraph("Skills: Python, FastAPI, MongoDB")
    doc.save(buf); buf.seek(0)
    return buf.read()

MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

@pytest_asyncio.fixture
async def ac():
    from backend.main import app
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as c:
        yield c


class TestResumeLimits:
    @pytest.mark.asyncio
    async def test_limits_returns_remaining(self, ac):
        patch_all_db(users_val=BASE_USER)
        r = await ac.get("/api/resume/limits",
                         headers={"Authorization": f"Bearer {make_jwt(UID)}"})
        assert r.status_code == 200
        assert "remaining" in r.json()

    @pytest.mark.asyncio
    async def test_unauthenticated_returns_401(self, ac):
        assert (await ac.get("/api/resume/limits")).status_code == 401


class TestResumeUpload:
    @pytest.mark.asyncio
    async def test_valid_docx_returns_feedback(self, ac):
        patch_all_db(users_val=BASE_USER)
        with patch("backend.controllers.resume_routes.get_feedback",
                   new_callable=AsyncMock, return_value=FAKE_FB):
            r = await ac.post("/api/resume/upload",
                files={"file": ("cv.docx", _docx(), MIME)},
                data={"job_title": "Software Engineer", "consent": "false"},
                headers={"Authorization": f"Bearer {make_jwt(UID)}"})
        assert r.status_code == 200
        assert "Score" in r.json()["feedback"]

    @pytest.mark.asyncio
    async def test_gibberish_job_title_returns_400(self, ac):
        patch_all_db(users_val=BASE_USER)
        r = await ac.post("/api/resume/upload",
            files={"file": ("cv.docx", _docx(), MIME)},
            data={"job_title": "asdfghjkl", "consent": "false"},
            headers={"Authorization": f"Bearer {make_jwt(UID)}"})
        assert r.status_code == 400

    @pytest.mark.asyncio
    async def test_unauthenticated_returns_401(self, ac):
        r = await ac.post("/api/resume/upload",
            files={"file": ("cv.docx", _docx(), MIME)},
            data={"job_title": "Dev", "consent": "false"})
        assert r.status_code == 401

    @pytest.mark.asyncio
    async def test_daily_limit_reached_returns_429(self, ac):
        patch_all_db(users_val={**BASE_USER, "daily_resume_count": 5,
                                "daily_reset_at": datetime.now(timezone.utc)})
        with patch("backend.controllers.resume_routes.get_feedback",
                   new_callable=AsyncMock, return_value=FAKE_FB):
            r = await ac.post("/api/resume/upload",
                files={"file": ("cv.docx", _docx(), MIME)},
                data={"job_title": "Software Engineer", "consent": "false"},
                headers={"Authorization": f"Bearer {make_jwt(UID)}"})
        assert r.status_code == 429

    @pytest.mark.asyncio
    async def test_admin_cannot_upload_returns_403(self, ac):
        patch_all_db(users_val={**BASE_USER, "role": "admin"})
        r = await ac.post("/api/resume/upload",
            files={"file": ("cv.docx", _docx(), MIME)},
            data={"job_title": "Engineer", "consent": "false"},
            headers={"Authorization": f"Bearer {make_jwt(UID, role='admin')}"})
        assert r.status_code == 403


class TestManualUpload:
    P = {"jobTitle": "Software Engineer",
         "experience": "3 years backend development with Python and FastAPI",
         "summary":    "Passionate developer building scalable web applications",
         "skills":     "Python, FastAPI, MongoDB, Docker, REST APIs",
         "achievement":"Reduced API response time by 40% through query optimisation",
         "consent": False}

    @pytest.mark.asyncio
    async def test_valid_input_returns_feedback(self, ac):
        patch_all_db(users_val=BASE_USER)
        with patch("backend.controllers.resume_routes.get_feedback",
                   new_callable=AsyncMock, return_value=FAKE_FB):
            r = await ac.post("/api/resume/manual-upload", json=self.P,
                              headers={"Authorization": f"Bearer {make_jwt(UID)}"})
        assert r.status_code == 200
        assert "feedback" in r.json()

    @pytest.mark.asyncio
    async def test_gibberish_job_title_rejected(self, ac):
        patch_all_db(users_val=BASE_USER)
        r = await ac.post("/api/resume/manual-upload",
                          json={**self.P, "jobTitle": "qwertyuiop"},
                          headers={"Authorization": f"Bearer {make_jwt(UID)}"})
        assert r.status_code == 400

    @pytest.mark.asyncio
    async def test_gibberish_skills_rejected(self, ac):
        patch_all_db(users_val=BASE_USER)
        r = await ac.post("/api/resume/manual-upload",
                          json={**self.P, "skills": "asdfghjklzxcvbnm"},
                          headers={"Authorization": f"Bearer {make_jwt(UID)}"})
        assert r.status_code == 400


class TestMyResumes:
    @pytest.mark.asyncio
    async def test_returns_list(self, ac):
        patch_all_db(users_val=BASE_USER,
                     resumes_items=[{"_id": "abc", "filename": "cv.pdf",
                                     "status": "pending", "created_at": None,
                                     "tags": [], "feedback": {}, "job_title": "Dev"}])
        r = await ac.get("/api/resume/my",
                         headers={"Authorization": f"Bearer {make_jwt(UID)}"})
        assert r.status_code == 200
        assert isinstance(r.json(), list)
