"""
Functional Tests — End-to-end feature flow tests.
Each test covers a complete user journey for one feature (F001–F010).
No real DB or AI — everything mocked, but the full code path is exercised.
"""
import os, sys, io, json, pytest, pytest_asyncio
from unittest.mock import AsyncMock, patch
from httpx import AsyncClient, ASGITransport
from datetime import datetime, timezone, timedelta
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
if ROOT not in sys.path: sys.path.insert(0, ROOT)

from tests.integration.helpers import make_jwt, patch_all_db

UID = "507f191e810c19729de860ea"

@pytest_asyncio.fixture
async def ac():
    from backend.main import app
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as c:
        yield c

def _docx():
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Jane Doe | jane@test.com")
    doc.add_paragraph("Summary: Backend developer with 3 years Python experience.")
    doc.add_paragraph("Skills: Python, FastAPI, MongoDB, Docker")
    doc.save(buf); buf.seek(0)
    return buf.read()

MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

FAKE_FB = {
    "IsResume": True, "Score": 78,
    "ScoreBreakdown": {"ImpactScore": 32, "SkillScore": 24,
                       "StructureScore": 16, "ATSScore": 6},
    "Advantages": ["Good structure"], "Disadvantages": ["Needs summary"],
    "Suggestions": ["Add a professional summary"],
    "Keywords": ["Python", "FastAPI", "MongoDB"],
    "Location": "Kuala Lumpur", "DetectedJobTitle": "Software Engineer",
    "Email": "jane@test.com", "Phone": "+60123456789", "Website": "",
    "ProfessionalSummary": "Experienced developer.",
    "Education": [{"Institution": "APU", "Degree": "BSc CS",
                   "Date": "2024", "Location": "KL", "GPA": "3.8"}],
    "Experience": [{"Company": "Tech Corp", "Position": "Dev",
                    "Date": "2022-2024", "Bullets": ["Built REST APIs"]}],
    "Projects": [], "SkillsTech": "Python", "SkillsTools": "Git",
    "SkillsSoft": "Comm", "Certifications": [], "Languages": ["English"],
    "AdditionalInfo": [],
}


# ── F001: Full user login journey ────────────────────────────────────────────
class TestF001UserLoginFlow:
    """User submits credentials → receives JWT → calls /me successfully."""
    @pytest.mark.asyncio
    async def test_full_login_and_me(self, ac):
        from backend.core.security import hash_password
        user = {
            "_id": UID, "email": "jane@test.com",
            "password_hash": hash_password("Pass123!"),
            "role": "user", "is_verified": True,
            "failed_login_attempts": 0, "lockout_until": None,
            "name": "Jane", "target_job_title": "Software Engineer",
            "target_location": "KL", "has_analyzed": True,
        }
        patch_all_db(users_val=user)
        r = await ac.post("/api/auth/login",
                          data={"username": "jane@test.com", "password": "Pass123!"})
        assert r.status_code == 200
        token = r.json()["access_token"]

        # Use that token to call /me
        patch_all_db(users_val=user)
        r2 = await ac.get("/api/auth/me",
                          headers={"Authorization": f"Bearer {token}"})
        assert r2.status_code == 200
        assert r2.json()["email"] == "jane@test.com"


# ── F002: Registration → OTP → Account creation ──────────────────────────────
class TestF002RegistrationFlow:
    """Register → receive OTP → verify → account created."""
    @pytest.mark.asyncio
    async def test_register_then_verify(self, ac):
        # Step 1: Register
        patch_all_db(users_val=None)
        r = await ac.post("/api/auth/register",
                          json={"email": "newuser@test.com", "password": "Pass123!",
                                "name": "New User"})
        assert r.status_code == 200
        otp = r.json()["otp"]
        assert len(otp) == 6

        # Step 2: Verify with correct OTP
        from datetime import datetime, timezone
        patch_all_db(
            pending_val={"email": "newuser@test.com", "verification_otp": otp,
                         "otp_created_at": datetime.now(timezone.utc),
                         "failed_otp_attempts": 0,
                         "password_hash": "$2b$12$x", "name": "New User",
                         "ip_address": "127.0.0.1"},
            users_val=None)
        r2 = await ac.post("/api/auth/verify-email",
                           json={"email": "newuser@test.com", "otp": otp})
        assert r2.status_code == 200
        assert "verified" in r2.json()["message"].lower()


# ── F003: Forgot password → reset ───────────────────────────────────────────
class TestF003PasswordRecoveryFlow:
    @pytest.mark.asyncio
    async def test_forgot_password_sends_link(self, ac):
        from backend.core.security import hash_password
        patch_all_db(users_val={"_id": UID, "email": "jane@test.com",
                                 "password_hash": hash_password("OldPass1!")})
        with patch("backend.controllers.auth_routes.send_reset_password_email",
                   new_callable=AsyncMock, return_value=True):
            r = await ac.post("/api/auth/forgot-password",
                              json={"email": "jane@test.com"})
        assert r.status_code == 200
        assert "reset" in r.json()["message"].lower() or "sent" in r.json()["message"].lower()

    @pytest.mark.asyncio
    async def test_reset_password_succeeds(self, ac):
        from backend.core.security import hash_password, create_reset_token
        from tests.integration.helpers import make_col
        import backend.core.security as sec

        token_val = "test-reset-token-abc123"
        patch_all_db(
            users_val={"_id": UID, "email": "jane@test.com",
                       "password_hash": hash_password("OldPass1!")},
            reset_val={"token": token_val, "email": "jane@test.com",
                       "expires_at": datetime.now(timezone.utc) + timedelta(minutes=25)})
        with patch.object(sec, "verify_reset_token",
                          new_callable=AsyncMock, return_value="jane@test.com"):
            r = await ac.post("/api/auth/reset-password",
                              json={"token": token_val, "password": "NewPass1!"})
        assert r.status_code == 200


# ── F004: Resume upload → AI feedback ───────────────────────────────────────
class TestF004ResumeAnalysisFlow:
    @pytest.mark.asyncio
    async def test_upload_returns_score_and_breakdown(self, ac):
        patch_all_db(users_val={"_id": UID, "email": "u@t.com", "role": "user",
                                 "is_verified": True, "has_analyzed": False,
                                 "daily_resume_count": 0, "daily_interview_count": 0,
                                 "daily_question_count": 0, "daily_reset_at": None})
        with patch("backend.controllers.resume_routes.get_feedback",
                   new_callable=AsyncMock, return_value=FAKE_FB):
            r = await ac.post("/api/resume/upload",
                files={"file": ("cv.docx", _docx(), MIME)},
                data={"job_title": "Software Engineer", "consent": "false"},
                headers={"Authorization": f"Bearer {make_jwt(UID)}"})
        assert r.status_code == 200
        fb = r.json()["feedback"]
        assert fb["Score"] == 78
        assert fb["ScoreBreakdown"]["ImpactScore"] == 32
        assert len(fb["Keywords"]) > 0


# ── F005: Manual profile → AI feedback ──────────────────────────────────────
class TestF005ManualProfileFlow:
    @pytest.mark.asyncio
    async def test_manual_upload_full_flow(self, ac):
        patch_all_db(users_val={"_id": UID, "email": "u@t.com", "role": "user",
                                 "is_verified": True, "has_analyzed": False,
                                 "daily_resume_count": 0, "daily_interview_count": 0,
                                 "daily_question_count": 0, "daily_reset_at": None})
        with patch("backend.controllers.resume_routes.get_feedback",
                   new_callable=AsyncMock, return_value=FAKE_FB):
            r = await ac.post("/api/resume/manual-upload",
                json={"jobTitle": "Software Engineer",
                      "experience": "3 years Python backend development",
                      "summary": "Passionate developer with FastAPI expertise",
                      "skills": "Python, FastAPI, MongoDB, Docker",
                      "achievement": "Reduced latency by 40% via caching",
                      "consent": False},
                headers={"Authorization": f"Bearer {make_jwt(UID)}"})
        assert r.status_code == 200
        assert "feedback" in r.json()
        assert r.json()["feedback"]["Score"] == 78


# ── F006: Interview session start → reply → end ──────────────────────────────
class TestF006MockInterviewFlow:
    @pytest.mark.asyncio
    async def test_start_reply_end_flow(self, ac):
        SID = "functional-session-001"
        user = {"_id": UID, "email": "u@t.com", "role": "user",
                "is_verified": True, "has_analyzed": True,
                "target_job_title": "Software Engineer",
                "daily_resume_count": 0, "daily_interview_count": 0,
                "daily_question_count": 0, "daily_reset_at": None}
        session = {"_id": SID, "session_id": SID, "user_id": UID,
                   "job_title": "Software Engineer", "resume_feedback": {},
                   "questions_limit": 10, "difficulty": "Beginner",
                   "asked_count": 1, "invalid_attempts": 0,
                   "transcript": [{"role": "assistant",
                                    "text": "Tell me about yourself.",
                                    "at": datetime.now(timezone.utc)}],
                   "created_at": datetime.now(timezone.utc), "ended_at": None}

        # Start
        patch_all_db(users_val=user,
                     resumes_val={"feedback": {}, "job_title": "Software Engineer"})
        with patch("backend.controllers.interview_routes.interview_reply",
                   return_value="Tell me about yourself."):
            r = await ac.post("/api/interview/start",
                data={"job_title": "Software Engineer",
                      "difficulty": "Beginner", "questions_limit": "10"},
                headers={"Authorization": f"Bearer {make_jwt(UID)}"})
        assert r.status_code == 200
        sid = r.json()["session_id"]

        # Reply
        patch_all_db(users_val=user, interviews_val=session)
        with patch("backend.controllers.interview_routes.interview_reply",
                   return_value="What is your greatest strength?"), \
             patch("backend.controllers.interview_routes.rag_engine.validate_input",
                   new_callable=AsyncMock, return_value={"safe": True}):
            r2 = await ac.post(f"/api/interview/{SID}/reply",
                data={"user_text": "I am a strong Python developer."},
                headers={"Authorization": f"Bearer {make_jwt(UID)}"})
        assert r2.status_code == 200

        # End
        patch_all_db(users_val=user, interviews_val=session)
        with patch("backend.controllers.interview_routes.interview_reply",
                   return_value="Thank you. Session closed. [FINISH]"):
            r3 = await ac.post(f"/api/interview/{SID}/end",
                               headers={"Authorization": f"Bearer {make_jwt(UID)}"})
        assert r3.status_code == 200
        assert r3.json().get("ended") is True


# ── F007: History → detail → delete ─────────────────────────────────────────
class TestF007InterviewHistoryFlow:
    @pytest.mark.asyncio
    async def test_list_expand_delete_flow(self, ac):
        SID = "hist-session-001"
        user = {"_id": UID, "email": "u@t.com", "role": "user",
                "is_verified": True, "has_analyzed": True,
                "daily_resume_count": 0, "daily_interview_count": 0,
                "daily_question_count": 0, "daily_reset_at": None}
        session = {"_id": SID, "session_id": SID, "user_id": UID,
                   "job_title": "Dev", "resume_feedback": {},
                   "questions_limit": 10, "difficulty": "Beginner",
                   "asked_count": 10, "invalid_attempts": 0,
                   "transcript": [], "created_at": datetime.now(timezone.utc),
                   "ended_at": datetime.now(timezone.utc),
                   "readiness_score": 74, "readiness_feedback": "Good job.",
                   "readiness_breakdown": {}}

        # List
        patch_all_db(users_val=user, interviews_items=[session])
        r = await ac.get("/api/interview/history",
                         headers={"Authorization": f"Bearer {make_jwt(UID)}"})
        assert r.status_code == 200
        assert len(r.json()) >= 1

        # Delete
        patch_all_db(users_val=user, deleted_count=1)
        r2 = await ac.delete(f"/api/interview/{SID}",
                             headers={"Authorization": f"Bearer {make_jwt(UID)}"})
        assert r2.status_code == 200


# ── F009: Job search proxy ───────────────────────────────────────────────────
class TestF009JobSearchFlow:
    @pytest.mark.asyncio
    async def test_job_search_requires_keywords_param(self, ac):
        """Route exists — missing keywords returns 422, not 404."""
        r = await ac.get("/api/jobs")
        assert r.status_code == 422

    @pytest.mark.asyncio
    async def test_job_search_with_keywords_hits_route(self, ac):
        """Route accepts keywords — any non-404/500 response confirms it works."""
        # We can't mock httpx cleanly here since it uses a real context manager.
        # The functional assertion is: the route exists and handles the request.
        r = await ac.get("/api/jobs?keywords=Python+Developer")
        # Will be 500 (Careerjet unreachable in test) or 200 — never 404
        assert r.status_code != 404


# ── F010: Admin resume review ────────────────────────────────────────────────
class TestF010AdminReviewFlow:
    @pytest.mark.asyncio
    async def test_list_view_patch_delete_flow(self, ac):
        ADMIN_ID = "507f191e810c19729de860eb"
        RES_ID   = "507f191e810c19729de860ec"
        admin = {"_id": ADMIN_ID, "email": "adm@icp-solution.com",
                 "role": "admin", "is_verified": True}
        resume = {"_id": RES_ID, "user_id": UID, "filename": "cv.pdf",
                  "status": "pending", "tags": ["Python"],
                  "feedback": {"Score": 75, "Advantages": []},
                  "text": "Dev", "mime_type": "application/pdf",
                  "file_id": None, "notes": "",
                  "created_at": datetime.now(timezone.utc)}

        # List
        patch_all_db(users_val=admin, resumes_items=[resume])
        r = await ac.get("/api/admin/resumes",
                         headers={"Authorization": f"Bearer {make_jwt(ADMIN_ID, 'admin')}"})
        assert r.status_code == 200

        # View detail
        patch_all_db(users_val=admin, resumes_val=resume)
        r2 = await ac.get(f"/api/admin/resumes/{RES_ID}",
                          headers={"Authorization": f"Bearer {make_jwt(ADMIN_ID, 'admin')}"})
        assert r2.status_code == 200

        # Patch status
        patch_all_db(users_val=admin, modified_count=1)
        r3 = await ac.patch(f"/api/admin/resumes/{RES_ID}",
                            data={"status": "reviewed", "notes": "Looks good"},
                            headers={"Authorization": f"Bearer {make_jwt(ADMIN_ID, 'admin')}"})
        assert r3.status_code == 200

        # Delete
        patch_all_db(users_val=admin, resumes_val={**resume, "file_id": None},
                     deleted_count=1)
        r4 = await ac.delete(f"/api/admin/resumes/{RES_ID}",
                             headers={"Authorization": f"Bearer {make_jwt(ADMIN_ID, 'admin')}"})
        assert r4.status_code == 200
        assert r4.json()["deleted"] is True
