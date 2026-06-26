"""
Unit Tests — backend/services/resume_parser.py
Tests: is_pdf, is_docx, clean_text, DOCX extraction
Uses real file I/O with temporary files; no DB or network needed.
"""
import os
import sys
import io
import tempfile
import pytest

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".."))
if ROOT not in sys.path:
    sys.path.insert(0, ROOT)

from backend.services.resume_parser import is_pdf, is_docx, clean_text, extract_resume_text
from docx import Document


def make_docx_file(text: str = "Software Engineer with Python skills") -> str:
    """Create a real temporary DOCX file and return its path."""
    buf = io.BytesIO()
    doc = Document()
    doc.add_heading("John Doe", level=1)
    doc.add_paragraph("Email: john@example.com")
    doc.add_paragraph("Summary")
    doc.add_paragraph(text)
    doc.add_paragraph("Education: BSc Computer Science, 2024")
    doc.add_paragraph("Skills: Python, FastAPI, MongoDB")
    doc.save(buf)
    buf.seek(0)

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.write(buf.read())
    tmp.close()
    return tmp.name


class TestFileTypeDetection:
    def test_pdf_detected(self):
        assert is_pdf("resume.pdf") is True
        assert is_pdf("resume.PDF") is True

    def test_docx_detected(self):
        assert is_docx("resume.docx") is True
        assert is_docx("resume.DOCX") is True

    def test_txt_not_pdf(self):
        assert is_pdf("resume.txt") is False

    def test_txt_not_docx(self):
        assert is_docx("resume.txt") is False


class TestCleanText:
    def test_removes_extra_spaces(self):
        result = clean_text("hello   world")
        assert "  " not in result

    def test_removes_extra_newlines(self):
        result = clean_text("line1\n\n\nline2")
        assert "\n\n" not in result

    def test_strips_whitespace(self):
        result = clean_text("   hello   ")
        assert result == "hello"

    def test_empty_returns_empty(self):
        assert clean_text("") == ""

    def test_none_returns_empty(self):
        assert clean_text(None) == ""


class TestDocxExtraction:
    def test_extract_docx_returns_text(self):
        path = make_docx_file("Python developer with 3 years of experience")
        try:
            text, mime, ocr_used = extract_resume_text(path)
            assert len(text) > 20
            assert "python" in text.lower() or "developer" in text.lower()
        finally:
            os.unlink(path)

    def test_extract_docx_mime_type(self):
        path = make_docx_file()
        try:
            _, mime, _ = extract_resume_text(path)
            assert "wordprocessingml" in mime or "docx" in mime.lower()
        finally:
            os.unlink(path)

    def test_extract_docx_ocr_false(self):
        path = make_docx_file()
        try:
            _, _, ocr_used = extract_resume_text(path)
            assert ocr_used is False
        finally:
            os.unlink(path)

    def test_unsupported_file_raises(self):
        tmp = tempfile.NamedTemporaryFile(suffix=".txt", delete=False)
        tmp.write(b"some text")
        tmp.close()
        try:
            with pytest.raises(ValueError, match="Unsupported file type"):
                extract_resume_text(tmp.name)
        finally:
            os.unlink(tmp.name)
